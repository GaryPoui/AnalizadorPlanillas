"""
Price Extractor API - Multi-Agent System
Orchestrates extraction of price data from various file formats
and maps them to the purchase price template format.
"""

import os
import io
import json
import base64
import re
import hashlib
import asyncio
import tempfile
import unicodedata
from pathlib import Path
from typing import Optional
from datetime import datetime

import httpx
import pandas as pd
import pdfplumber
import pytesseract
from dotenv import load_dotenv
from docx import Document as DocxDocument
from PIL import Image
from markitdown import MarkItDown
from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, StreamingResponse
from pydantic import BaseModel

app = FastAPI(title="Price Extractor API", version="1.0.0")

# Load local env files for non-Docker runs.
API_DIR = Path(__file__).resolve().parent
load_dotenv(API_DIR / ".env", override=True)
load_dotenv(API_DIR.parent / ".env", override=True)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Anthropic (Claude) configuration
CLAUDE_MODEL = os.getenv("CLAUDE_MODEL", "claude-haiku-4-5")
MARKITDOWN = MarkItDown()


def detect_extension(filename: str, file_bytes: bytes) -> str:
    """Detect extension from filename, with content-based fallback for malformed multipart names."""
    ext = Path(filename or "").suffix.lower().strip()
    if ext:
        return ext

    # Content signature fallback.
    sig = file_bytes[:16]
    if sig.startswith(b"%PDF"):
        return ".pdf"
    if sig.startswith(b"PK\x03\x04"):
        # Could be xlsx/xlsm/docx; API supports xlsx/xlsm here.
        return ".xlsx"
    if sig.startswith(b"\xFF\xD8\xFF"):
        return ".jpg"
    if sig.startswith(b"\x89PNG\r\n\x1a\n"):
        return ".png"
    if sig.startswith(b"RIFF") and b"WEBP" in file_bytes[:32]:
        return ".webp"
    if sig.startswith(b"BM"):
        return ".bmp"

    # Text heuristics: csv default if separators are present.
    try:
        sample = file_bytes[:4096].decode("utf-8", errors="ignore")
        if sample.count(",") or sample.count(";") or sample.count("\t"):
            return ".csv"
    except Exception:
        pass

    return ""


async def claude_chat(messages: list, system: str = "", max_tokens: int = 8000) -> str:
    """Send a chat request to Anthropic Claude API and return the response text."""
    api_key = os.getenv("ANTHROPIC_API_KEY", "").strip()
    if not api_key:
        raise HTTPException(
            status_code=500,
            detail=(
                "ANTHROPIC_API_KEY is not configured. "
                "Set it in your environment or .env file."
            ),
        )

    headers = {
        "x-api-key": api_key,
        "anthropic-version": "2023-06-01",
        "content-type": "application/json",
    }
    payload = {
        "model": CLAUDE_MODEL,
        "max_tokens": max_tokens,
        "messages": messages,
    }
    if system:
        payload["system"] = system

    async with httpx.AsyncClient(timeout=300.0) as client:
        try:
            resp = await client.post(
                "https://api.anthropic.com/v1/messages",
                headers=headers,
                json=payload,
            )
            resp.raise_for_status()
        except httpx.HTTPStatusError as e:
            detail = e.response.text[:600] if e.response is not None else str(e)
            raise HTTPException(
                status_code=502,
                detail=f"Anthropic API error ({e.response.status_code}): {detail}",
            )
        except httpx.HTTPError as e:
            raise HTTPException(
                status_code=502,
                detail=f"Anthropic connection error: {str(e)}",
            )

        data = resp.json()
        content_blocks = data.get("content", [])
        texts = []
        for block in content_blocks:
            if isinstance(block, dict) and block.get("text"):
                texts.append(block["text"])

        if texts:
            return "\n".join(texts)

        raise HTTPException(
            status_code=502,
            detail=(
                "Anthropic response did not include text content. "
                f"Raw: {json.dumps(data)[:600]}"
            ),
        )

# Template columns matching Plantilla_Precios_Compras
TEMPLATE_COLUMNS = [
    "Cód. Artículo",
    "Descripción artículo",
    "Descripción adicional artículo",
    "Sinónimo",
    "Cód. Lista",
    "Desc. Lista",
    "Moneda",
    "Unidad",
    "Precio",
    "Bonif.",
    "Fecha vigencia desde",
    "Fecha vigencia hasta",
]

TRANSFORM_CACHE_MAX_ITEMS = max(int(os.getenv("TRANSFORM_CACHE_MAX_ITEMS", "512")), 32)
TRANSFORM_CACHE: dict[str, dict] = {}
AI_COMPLEMENT_ALL_CHUNKS = os.getenv("AI_COMPLEMENT_ALL_CHUNKS", "0") == "1"
PDF_MIN_EXPECTED_ROWS = max(int(os.getenv("PDF_MIN_EXPECTED_ROWS", "80")), 10)
PDF_RECOVERY_MAX_CHUNKS = max(int(os.getenv("PDF_RECOVERY_MAX_CHUNKS", "12")), 1)
# Hybrid mode: per-page (PDF) or per-chunk (XLS) Claude pass after heuristic
HYBRID_EXTRACTION = os.getenv("HYBRID_EXTRACTION", "1") == "1"
HYBRID_MAX_TOKENS = max(int(os.getenv("HYBRID_MAX_TOKENS", "6000")), 2000)
HYBRID_XLS_CHUNK_CHARS = max(int(os.getenv("HYBRID_XLS_CHUNK_CHARS", "25000")), 5000)

# Domain-aware system prompt built from sessions 1-9 experience
_HYBRID_SYSTEM_PROMPT = (
    "You extract product entries from Argentine electrical/industrial supplier price lists.\n"
    "Return ONLY a JSON array. Each object must have:\n"
    '  "code": product code — 4-5 digit number OR alphanumeric (e.g. SCA-10, UCA-16, HM-12CB, A2, B3)\n'
    '  "desc": full product description including model name and specifications\n'
    '  "price": unit price as a plain decimal number (no currency symbols, no spaces)\n'
    "\n"
    "Rules:\n"
    "- Extract ALL products visible on the page/section\n"
    "- When a page has TWO side-by-side product columns, extract from BOTH columns\n"
    "- For paired rows (two products on one line), produce two separate objects\n"
    "- Convert Argentine price format: '1.234,56' → 1234.56; '1 234,56' → 1234.56\n"
    "- Skip: page headers, section/category titles, subtotals, empty rows\n"
    "- Known supplier formats: LCT (4-5 digit codes), N°95 (4-5 digit), MICROCONTROL (alphanumeric)\n"
    "- Return [] if the page/section has no product rows"
)

# Codes: alphanum with dashes OR pure 4-5 digit (NOT 6+ digit barcodes/prices)
CODE_TOKEN_RE = r"(?:[A-Z]{1,8}-\d{1,5}(?:-\d{1,5})?|[A-Z]{1,5}\d{2,6}|\d{4,5})"
PRICE_TOKEN_RE = r"(?:\$\s*)?\d{1,3}(?:\.\d{3})*(?:,\d{2})|(?:\$\s*)?\d{4,7}"
CODE_PRICE_RE = re.compile(rf"({CODE_TOKEN_RE})\s+({PRICE_TOKEN_RE})")
CODE_PRICE_ANYWHERE_RE = re.compile(rf"({CODE_TOKEN_RE}).{{0,80}}?({PRICE_TOKEN_RE})")
CODE_ONLY_RE = re.compile(rf"\b({CODE_TOKEN_RE})\b", flags=re.IGNORECASE)
PRICE_ONLY_RE = re.compile(rf"\b({PRICE_TOKEN_RE})\b")

# Words that are never valid product codes (headers, materials, common words)
_CODE_BLACKLIST = frozenset({
    "CODIGO", "REFERENCIA", "PRECIO", "DESCRIPCION", "MODELO", "DETALLE",
    "ALUMINIO", "COBRE", "BRONCE", "HIERRO", "ACERO", "PVC", "POLIETILENO",
    "SIN", "POR", "CON", "DEL", "LOS", "LAS",
    "ISO9001", "QM15", "IP23", "IP24", "IEC",
    # LCT tool model names (captured as codes by word-coord; real code is numeric 4xxx)
    "PH-6", "HT-240U", "HT-240C", "HM-12CB", "CO-12CB", "CO-400", "CO-630",
    "CO-IP", "HH-IP", "HEC-240", "HEC-400", "MHEC-240", "MHEC-240S",
    # Norm/spec labels that appear in product table headers
    "IEC-6", "APTAS", "LY-468",
})

# Price pattern: $? optional, also handle OCR-spaced digits
ITEM_LINE_RE = re.compile(r"^\s*(?P<item_code>\d{4,5})\b.*?(?P<price>\$?\s*\d[\d\s]*\d?[,.]\d{2}|\$\s*\d{4,7})")
NUMERIC_ITEM_PROFILE_RE = re.compile(r"^\s*\d{4,5}\b.*?\$?\s*\d[\d\s]*\d?[,.]\d{2}", re.IGNORECASE)
# Multi-product line: extract ALL code+price pairs from same line
MULTI_ITEM_LINE_RE = re.compile(r"(\d{4,5})\s+[A-Z][^$\d]*\$\s*(\d[\d\s]*\d?[,.]\d{2})", re.IGNORECASE)
# Transposed table: codes on one line, models on next line, prices further down
TRANSPOSED_CODES_RE = re.compile(r"^(\d{4,5}(?:\s+\d{4,5}){2,})\s*$")
CANDIDATE_LINE_RE = re.compile(
    rf"({CODE_TOKEN_RE}|\d{{1,3}}(?:[\.,]\d{{3}})*[\.,]\d{{2}}|\d{{4,9}}|USD|U\$S|\$|bonif|vigencia)",
    flags=re.IGNORECASE,
)


def _cache_get(cache_key: str) -> Optional[dict]:
    cached = TRANSFORM_CACHE.get(cache_key)
    if cached is None:
        return None
    # Keep recently used entries at the end.
    TRANSFORM_CACHE.pop(cache_key, None)
    TRANSFORM_CACHE[cache_key] = cached
    return cached


def _cache_set(cache_key: str, value: dict) -> None:
    if cache_key in TRANSFORM_CACHE:
        TRANSFORM_CACHE.pop(cache_key, None)
    TRANSFORM_CACHE[cache_key] = value
    while len(TRANSFORM_CACHE) > TRANSFORM_CACHE_MAX_ITEMS:
        oldest_key = next(iter(TRANSFORM_CACHE))
        TRANSFORM_CACHE.pop(oldest_key, None)


def _normalize_price_token(price_text: str) -> str:
    normalized = str(price_text).strip().replace("$", "").replace(" ", "")
    # Handle OCR-spaced digits like "7 0 0 1 ,66" → "7001,66"
    if re.search(r'\d\s+\d', str(price_text)):
        normalized = re.sub(r'(?<=\d)\s+(?=\d)', '', str(price_text)).replace("$", "").replace(" ", "")
    if "," in normalized and "." in normalized:
        if normalized.rfind(",") > normalized.rfind("."):
            normalized = normalized.replace(".", "").replace(",", ".")
        else:
            normalized = normalized.replace(",", "")
    elif "," in normalized:
        normalized = normalized.replace(".", "").replace(",", ".")
    normalized = normalized.strip()
    return normalized


def _infer_currency(text: str) -> str:
    upper = text.upper()
    if "U$S" in upper or "USD" in upper or "US$" in upper:
        return "USD"
    if "EUR" in upper:
        return "EUR"
    return "ARS"


def _extract_list_header(text: str) -> tuple[str, str]:
    match = re.search(r"\blista\s*(?:n[°oº]\s*)?([A-Z0-9.-]{1,20})", text, flags=re.IGNORECASE)
    if not match:
        return "", ""
    list_code = match.group(1).strip().upper()
    list_desc = f"Lista {list_code}" if list_code else ""
    return list_code, list_desc


def _empty_template_row() -> dict:
    return {col: "" for col in TEMPLATE_COLUMNS}


def _row_key(row: dict) -> tuple[str, str]:
    return (
        str(row.get("Cód. Artículo", "")).strip().lower(),
        str(row.get("Precio", "")).strip().lower(),
    )


def _row_non_empty_score(row: dict) -> int:
    return sum(1 for v in row.values() if str(v).strip())


def _format_key_samples(keys: set[tuple[str, str]], limit: int = 20) -> list[dict]:
    out = []
    for code, price in sorted(keys)[:limit]:
        out.append({"Cód. Artículo": code, "Precio": price})
    return out


def build_candidate_blocks(raw_text: str, max_chunk_chars: int = 6000) -> list[str]:
    if not raw_text.strip():
        return []

    if "=== PAGE" in raw_text:
        pages = re.split(r"(?=^=== PAGE\s+\d+\s+===)", raw_text, flags=re.MULTILINE)
    else:
        pages = [raw_text]

    reduced_pages = []
    for page in pages:
        if not page.strip():
            continue
        lines = [ln for ln in page.splitlines() if ln.strip()]
        if not lines:
            continue

        selected_idxs = set()
        for i, line in enumerate(lines):
            if CANDIDATE_LINE_RE.search(line):
                # Keep a wider context window to avoid losing rows broken across lines.
                selected_idxs.update({i - 2, i - 1, i, i + 1, i + 2})

        if not selected_idxs:
            continue

        selected = []
        for idx in sorted(x for x in selected_idxs if 0 <= x < len(lines)):
            selected.append(lines[idx])

        # If filtering was too aggressive for a page, keep the full page content.
        if len(selected) < 8:
            selected = lines

        reduced = "\n".join(selected).strip()
        if reduced:
            reduced_pages.append(reduced)

    if not reduced_pages:
        reduced_pages = [raw_text]

    chunks = []
    current = ""
    for page_text in reduced_pages:
        if len(current) + len(page_text) + 1 > max_chunk_chars and current:
            chunks.append(current)
            current = page_text
        else:
            current = f"{current}\n{page_text}".strip()

    if current.strip():
        chunks.append(current)

    return [c for c in chunks if c.strip()]


def build_full_chunks(raw_text: str, max_chunk_chars: int = 4500) -> list[str]:
    """Split full text without aggressive filtering for recovery passes."""
    if not raw_text.strip():
        return []

    if "=== PAGE" in raw_text:
        pages = re.split(r"(?=^=== PAGE\s+\d+\s+===)", raw_text, flags=re.MULTILINE)
    else:
        pages = [raw_text]

    chunks = []
    current = ""
    for page_text in pages:
        page_text = page_text.strip()
        if not page_text:
            continue
        if len(current) + len(page_text) + 1 > max_chunk_chars and current:
            chunks.append(current)
            current = page_text
        else:
            current = f"{current}\n{page_text}".strip()

    if current.strip():
        chunks.append(current)

    return [c for c in chunks if c.strip()]


def heuristic_extract_rows(
    text: str,
    list_code: str = "",
    list_desc: str = "",
    default_currency: str = "ARS",
) -> list[dict]:
    rows = []

    # High precision path for catalogs where each item starts with numeric code.
    for line in text.splitlines():
        line_clean = line.strip()
        if not line_clean:
            continue
        m = ITEM_LINE_RE.search(line_clean)
        if not m:
            continue

        code_norm = m.group("item_code").strip().upper()
        price_norm = _normalize_price_token(m.group("price"))
        try:
            if float(price_norm) <= 3:
                continue
        except Exception:
            continue

        row = _empty_template_row()
        row["Cód. Artículo"] = code_norm
        row["Precio"] = price_norm

        # Keep only the core description segment between code and price to avoid noise.
        line_wo_code = re.sub(rf"^\s*{re.escape(code_norm)}\b", "", line_clean).strip()
        line_wo_price = line_wo_code.replace(m.group("price"), "").strip()
        desc = re.sub(r"\s{2,}", " ", line_wo_price).strip(" |;\t")
        row["Descripción artículo"] = desc if desc else code_norm
        row["Cód. Lista"] = list_code
        row["Desc. Lista"] = list_desc
        row["Moneda"] = default_currency
        row["Unidad"] = "Un"
        rows.append(row)

    if rows:
        # Preserve appearance order and dedupe by code+price.
        deduped = []
        seen = set()
        for row in rows:
            key = _row_key(row)
            if key in seen:
                continue
            seen.add(key)
            deduped.append(row)
        return deduped

    # Generic fallback path.
    for line in text.splitlines():
        line_clean = line.strip()
        if not line_clean:
            continue
        matches = CODE_PRICE_RE.findall(line_clean)
        if not matches:
            matches = CODE_PRICE_ANYWHERE_RE.findall(line_clean)
        if not matches:
            continue

        for code, raw_price in matches:
            row = _empty_template_row()
            code_norm = code.strip().upper()
            row["Cód. Artículo"] = code_norm
            row["Precio"] = _normalize_price_token(raw_price)

            # Skip clearly invalid low-value integer prices (noise near headers).
            try:
                if float(row["Precio"]) <= 3:
                    continue
            except Exception:
                continue

            # Keep a usable description while avoiding expensive IA for obvious rows.
            line_without_code_price = re.sub(rf"\b{re.escape(code)}\b", "", line_clean, flags=re.IGNORECASE)
            line_without_code_price = line_without_code_price.replace(raw_price, "")
            desc = re.sub(r"\s{2,}", " ", line_without_code_price).strip(" |;-\t")
            row["Descripción artículo"] = desc if desc else code.strip()

            row["Cód. Lista"] = list_code
            row["Desc. Lista"] = list_desc
            row["Moneda"] = default_currency
            row["Unidad"] = "Un"

            rows.append(row)

    deduped = []
    seen = set()
    for row in rows:
        key = (
            row.get("Cód. Artículo", "").strip().lower(),
            row.get("Precio", "").strip().lower(),
        )
        if key in seen:
            continue
        seen.add(key)
        deduped.append(row)
    return deduped


def heuristic_extract_rows_blockwise(
    text: str,
    list_code: str = "",
    list_desc: str = "",
    default_currency: str = "ARS",
) -> list[dict]:
    """
    Fallback heuristic for difficult layouts: pair code lines with nearby price lines.
    """
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    if not lines:
        return []

    rows = []
    for i, line in enumerate(lines):
        code_match = CODE_ONLY_RE.search(line)
        if not code_match:
            continue

        code = code_match.group(1).strip().upper()

        # Search nearest price within +/- 2 lines.
        price_val = ""
        best_dist = 999
        for j in range(max(0, i - 2), min(len(lines), i + 3)):
            for pm in PRICE_ONLY_RE.findall(lines[j]):
                norm = _normalize_price_token(pm)
                try:
                    if float(norm) <= 3:
                        continue
                except Exception:
                    continue

                dist = abs(j - i)
                if dist < best_dist:
                    best_dist = dist
                    price_val = norm

        if not price_val:
            continue

        row = _empty_template_row()
        row["Cód. Artículo"] = code
        row["Precio"] = price_val
        row["Descripción artículo"] = code
        row["Cód. Lista"] = list_code
        row["Desc. Lista"] = list_desc
        row["Moneda"] = default_currency
        row["Unidad"] = "Un"
        rows.append(row)

    deduped = []
    seen = set()
    for row in rows:
        key = _row_key(row)
        if key in seen or not any(key):
            continue
        seen.add(key)
        deduped.append(row)
    return deduped


def heuristic_extract_rows_numeric_profile(
    text: str,
    list_code: str = "",
    list_desc: str = "",
    default_currency: str = "ARS",
) -> list[dict]:
    """
    Strict profile for catalogs where each product row starts with numeric code
    and contains an inline ARS price (e.g., LCT-like lists).
    Handles both single-product lines and multi-product lines.
    """
    rows = []
    lines = text.splitlines()
    # First pass: multi-product lines (e.g. "2170 SCA10 ... $ 896,85 2200 UCA10 ... $ 708,30")
    for line in lines:
        line_clean = line.strip()
        if not line_clean:
            continue
        multi_matches = MULTI_ITEM_LINE_RE.findall(line_clean)
        if len(multi_matches) >= 2:
            for code, raw_price in multi_matches:
                price = _normalize_price_token(raw_price)
                try:
                    if float(price) <= 3:
                        continue
                except Exception:
                    continue
                r = _empty_template_row()
                r["Cód. Artículo"] = code.strip()
                r["Precio"] = price
                r["Descripción artículo"] = code.strip()
                r["Cód. Lista"] = list_code
                r["Desc. Lista"] = list_desc
                r["Moneda"] = default_currency
                r["Unidad"] = "Un"
                rows.append(r)
            continue
        # Standard single-product line
        m = ITEM_LINE_RE.search(line_clean)
        if not m:
            continue
        code = m.group("item_code").strip()
        if not code:
            continue
        price = _normalize_price_token(m.group("price"))
        try:
            if float(price) <= 3:
                continue
        except Exception:
            continue
        r = _empty_template_row()
        r["Cód. Artículo"] = code
        r["Precio"] = price
        desc = re.sub(rf"^\s*{re.escape(code)}\b", "", line_clean).strip()
        desc = desc.replace(m.group("price"), "").strip(" |;\t")
        desc = re.sub(r"\s{2,}", " ", desc)
        r["Descripción artículo"] = desc if desc else code
        r["Cód. Lista"] = list_code
        r["Desc. Lista"] = list_desc
        r["Moneda"] = default_currency
        r["Unidad"] = "Un"
        rows.append(r)

    # Second pass: transposed tables (codes on one row, models next row, prices somewhere below)
    # e.g. "3000 3001 3002 ..." then "A2 A3 A4 ..." then ... then "$ 1234 $ 5678 ..."
    i = 0
    while i < len(lines):
        line_clean = lines[i].strip()
        tc = TRANSPOSED_CODES_RE.match(line_clean)
        if tc:
            codes_in_row = tc.group(1).split()
            # Scan next ~15 lines for matching count of prices
            price_re_scan = re.compile(r"\$\s*[\d\s]*[\d],\d{2}")
            for j in range(i + 1, min(i + 16, len(lines))):
                prices_found = price_re_scan.findall(lines[j])
                if len(prices_found) >= len(codes_in_row):
                    for k, (c, p) in enumerate(zip(codes_in_row, prices_found)):
                        price_n = _normalize_price_token(p)
                        try:
                            if float(price_n) <= 3:
                                continue
                        except Exception:
                            continue
                        r = _empty_template_row()
                        r["Cód. Artículo"] = c
                        r["Precio"] = price_n
                        r["Descripción artículo"] = c
                        r["Cód. Lista"] = list_code
                        r["Desc. Lista"] = list_desc
                        r["Moneda"] = default_currency
                        r["Unidad"] = "Un"
                        rows.append(r)
                    break
        i += 1

    deduped = []
    seen = set()
    for row in rows:
        key = _row_key(row)
        if key in seen:
            continue
        seen.add(key)
        deduped.append(row)
    return deduped


def _is_alphanumeric_code(code: str) -> bool:
    code = str(code or "").strip()
    return bool(code) and any(ch.isalpha() for ch in code) and any(ch.isdigit() for ch in code)


def _try_ocr(file_bytes: bytes, num_pages: int) -> str:
    """OCR fallback for scanned PDFs using pytesseract (BUG-2)."""
    try:
        import pytesseract
        tess_cmd = os.getenv("TESSERACT_CMD", "").strip()
        if tess_cmd:
            pytesseract.pytesseract.tesseract_cmd = tess_cmd
        langs = "spa+eng"
        pages_text: list[str] = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for i, page in enumerate(pdf.pages[:num_pages]):
                try:
                    img = page.to_image(resolution=200).original
                    text = pytesseract.image_to_string(img, lang=langs).strip()
                    if text:
                        pages_text.append(f"=== PAGE {i + 1} ===\n{text}")
                except Exception:
                    continue
        return "\n".join(pages_text).strip()
    except Exception:
        return ""


def _extract_rows_from_word_coords(
    page,
    list_code: str = "",
    list_desc: str = "",
    default_currency: str = "ARS",
) -> list[dict]:
    """Supplement table detection using word bounding-box coordinates (BUG-3)."""
    try:
        words = page.extract_words(x_tolerance=3, y_tolerance=3, keep_blank_chars=False)
    except Exception:
        return []
    if not words:
        return []

    lines: dict[int, list] = {}
    for w in words:
        bucket = round(float(w.get("top", 0)) / 5) * 5
        lines.setdefault(bucket, []).append(w)

    price_re = re.compile(r"^\$?\s*\d{1,3}(?:\.\d{3})*[,\.]\d{2}$")
    rows = []
    for bucket in sorted(lines):
        line_words = sorted(lines[bucket], key=lambda w: float(w.get("x0", 0)))
        texts = [w["text"].strip() for w in line_words if w["text"].strip()]
        if len(texts) < 2:
            continue
        code = next((t for t in texts if _is_valid_product_code(t)), None)
        if not code:
            continue
        price_str = ""
        for t in reversed(texts):
            if price_re.match(t):
                candidate = _normalize_price_token(t)
                try:
                    if float(candidate) > 3:
                        price_str = candidate
                        break
                except Exception:
                    pass
        if not price_str:
            continue
        code_idx = next((i for i, t in enumerate(texts) if t == code), 0)
        rev_price_idx = next((i for i, t in enumerate(reversed(texts)) if price_re.match(t)), 0)
        price_idx = len(texts) - 1 - rev_price_idx
        desc_parts = [
            t for t in texts[code_idx + 1:price_idx]
            if not _is_valid_product_code(t) and not price_re.match(t)
        ]
        r = _empty_template_row()
        r["Cód. Artículo"] = code
        r["Precio"] = price_str
        r["Descripción artículo"] = " ".join(desc_parts).strip() or code
        r["Cód. Lista"] = list_code
        r["Desc. Lista"] = list_desc
        r["Moneda"] = default_currency
        r["Unidad"] = "Un"
        rows.append(r)
    return rows


def _extract_pdf_sync(file_bytes: bytes) -> dict:
    """All synchronous PDF work in one call — enables asyncio.to_thread() timeout (BUG-4)."""
    md_text = ""
    pdfplumber_text = ""
    pdf_raw_tables: list = []
    pdf_word_rows: list = []
    pages = 0

    try:
        result = MARKITDOWN.convert(io.BytesIO(file_bytes))
        md_text = (result.markdown or "").strip()
    except Exception:
        pass

    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            pages = len(pdf.pages)
            pages_text = []
            for i, page in enumerate(pdf.pages):
                text = (page.extract_text() or "").strip()
                tables = page.extract_tables() or []
                if tables:
                    pdf_raw_tables.append((i + 1, tables))
                table_lines = []
                for table in tables:
                    for row in table or []:
                        if row:
                            table_lines.append(" | ".join(str(c or "") for c in row))
                block = f"=== PAGE {i+1} ===\n{text}"
                if table_lines:
                    block += "\n" + "\n".join(table_lines)
                pages_text.append(block)
                pdf_word_rows.extend(_extract_rows_from_word_coords(page))
            pdfplumber_text = "\n".join(pages_text).strip()
    except Exception:
        pass

    sources = []
    if md_text:
        sources.append(f"=== SOURCE: MARKITDOWN ===\n=== PAGE 1 ===\n{md_text}")
    if pdfplumber_text:
        sources.append(f"=== SOURCE: PDFPLUMBER ===\n{pdfplumber_text}")
    raw_text = "\n\n".join(sources).strip()

    if md_text and pdfplumber_text:
        extraction_method = "pdf_dual"
    elif md_text:
        extraction_method = "pdf_markitdown_only"
    else:
        extraction_method = "pdf_pdfplumber_only"

    num_pages = max(pages, 1)
    if len(raw_text) / num_pages < 1500:
        ocr_text = _try_ocr(file_bytes, num_pages)
        if ocr_text:
            raw_text = ocr_text
            extraction_method = "pdf_ocr"

    return {
        "raw_text": raw_text,
        "pdf_raw_tables": pdf_raw_tables,
        "pdf_word_rows": pdf_word_rows,
        "pdf_pages": [p.strip() for p in pages_text],  # per-page text for hybrid pass
        "pages": pages,
        "md_chars": len(md_text),
        "plumber_chars": len(pdfplumber_text),
        "extraction_method": extraction_method,
    }


def _is_valid_product_code(code: str) -> bool:
    """True if code looks like a real product code (not a 6+ digit barcode or price)."""
    code = str(code or "").strip().upper()
    if not code:
        return False
    if code in _CODE_BLACKLIST:
        return False
    if code.isdigit():
        return 4 <= len(code) <= 5
    return bool(re.match(r'^[A-Z][A-Z0-9\-/]{1,15}$', code))


def _detect_table_column_roles(rows: list[list[str]]) -> dict[str, int]:
    """Detect which column in a pdfplumber table is code / price / description."""
    if not rows:
        return {"code": 0, "price": -1, "description": -1}
    num_cols = max((len(r) for r in rows), default=1)
    code_score = [0] * num_cols
    price_score = [0] * num_cols
    desc_score = [0] * num_cols

    header = [str(c or "").lower().strip() for c in rows[0]]
    for i, h in enumerate(header[:num_cols]):
        if any(w in h for w in ["cód", "cod", "código", "codigo", "ref", "item", "art"]):
            code_score[i] += 15
        if any(w in h for w in ["desc", "detalle", "modelo", "nombre", "product"]):
            desc_score[i] += 15
        if any(w in h for w in ["precio", "price", "$", "valor", "unitario"]):
            price_score[i] += 15

    price_cell_re = re.compile(r"^\$?\s*\d{1,3}(?:\.\d{3})*,\d{2}$")
    for row in rows[1:31]:
        for i, cell in enumerate(row[:num_cols]):
            cell_s = str(cell or "").strip()
            if not cell_s:
                continue
            if _is_valid_product_code(cell_s):
                code_score[i] += 4
            if price_cell_re.match(cell_s):
                price_score[i] += 5
            if len(cell_s) > 15:
                desc_score[i] += 2

    code_col = code_score.index(max(code_score)) if max(code_score) > 0 else 0
    price_candidates = sorted(range(num_cols), key=lambda x: price_score[x], reverse=True)
    price_col = next((c for c in price_candidates if c != code_col and price_score[c] > 0), -1)
    desc_candidates = [i for i in range(num_cols) if i != code_col and i != price_col]
    desc_col = max(desc_candidates, key=lambda x: desc_score[x]) if desc_candidates else -1

    # Detect extra code+price pairs for multi-column layouts (multiple products per row)
    extra_pairs: list[tuple[int, int]] = []
    used_cols: set[int] = {code_col, price_col}
    for i in range(num_cols):
        if i in used_cols or code_score[i] < 2:
            continue
        for j in (i + 1, i + 2, i - 1):
            if j < 0 or j >= num_cols or j in used_cols or price_score[j] < 2:
                continue
            extra_pairs.append((i, j))
            used_cols.update({i, j})
            break
    return {"code": code_col, "price": price_col, "description": desc_col, "extra_pairs": extra_pairs}


def extract_rows_from_pdf_tables(
    pdf_raw_tables: list,
    list_code: str = "",
    list_desc: str = "",
    default_currency: str = "ARS",
) -> list[dict]:
    """Extract product rows from pdfplumber raw table data using column role detection."""
    price_scan_re = re.compile(r"\$?\s*\d{1,3}(?:\.\d{3})*,\d{2}")
    all_rows: list[dict] = []

    for _page_num, tables in pdf_raw_tables:
        for table in tables:
            if not table or len(table) < 2:
                continue
            norm: list[list[str]] = []
            for row in table:
                cells = [str(c or "").strip() if c is not None else "" for c in row]
                if any(cells):
                    norm.append(cells)
            if len(norm) < 2:
                continue

            # Detect transposed single-column tables: CODE / MODEL / DESC / $ PRICE
            # Common in tool/product catalogs where each column = one product.
            num_cols = max(len(r) for r in norm) if norm else 1
            if num_cols == 1 and len(norm) >= 2:
                flat = [r[0] for r in norm if r]
                # Find a numeric code at the start and a price at the end
                code_candidate = flat[0] if flat else ""
                if _is_valid_product_code(code_candidate) and code_candidate.isdigit():
                    price_candidate = ""
                    desc_parts = []
                    for cell in flat[1:]:
                        m_price = price_scan_re.search(cell)
                        if m_price and not price_candidate:
                            price_candidate = _normalize_price_token(m_price.group())
                            try:
                                if float(price_candidate) <= 3:
                                    price_candidate = ""
                            except Exception:
                                price_candidate = ""
                        elif not m_price:
                            desc_parts.append(cell)
                    if price_candidate:
                        r = _empty_template_row()
                        r["Cód. Artículo"] = code_candidate
                        r["Precio"] = price_candidate
                        r["Descripción artículo"] = " ".join(desc_parts).strip() or code_candidate
                        r["Cód. Lista"] = list_code
                        r["Desc. Lista"] = list_desc
                        r["Moneda"] = default_currency
                        r["Unidad"] = "Un"
                        all_rows.append(r)
                        continue  # skip standard column-role detection for this table

            cols = _detect_table_column_roles(norm)
            code_col = cols["code"]
            price_col = cols["price"]
            desc_col = cols["description"]
            extra_pairs = cols.get("extra_pairs", [])

            first_cell = norm[0][code_col] if code_col < len(norm[0]) else ""
            data_start = 0 if _is_valid_product_code(first_cell) else 1

            for row in norm[data_start:]:
                if len(row) <= code_col:
                    continue
                code = str(row[code_col] if code_col < len(row) else "").strip().upper()
                if not code or not _is_valid_product_code(code):
                    continue

                price_str = ""
                if price_col >= 0 and price_col < len(row):
                    price_str = _normalize_price_token(str(row[price_col] or ""))
                    try:
                        if float(price_str) <= 3:
                            price_str = ""
                    except Exception:
                        price_str = ""

                if not price_str:
                    for i, cell in enumerate(row):
                        if i == code_col:
                            continue
                        m = price_scan_re.search(str(cell or ""))
                        if m:
                            candidate = _normalize_price_token(m.group())
                            try:
                                if float(candidate) > 3:
                                    price_str = candidate
                                    break
                            except Exception:
                                pass

                if not price_str:
                    continue

                desc = ""
                if desc_col >= 0 and desc_col < len(row):
                    desc = str(row[desc_col] or "").strip()
                if not desc:
                    parts = [
                        str(row[i] or "").strip()
                        for i in range(len(row))
                        if i not in (code_col, price_col) and str(row[i] or "").strip()
                    ]
                    desc = " ".join(parts[:2]).strip()

                r = _empty_template_row()
                r["Cód. Artículo"] = code
                r["Precio"] = price_str
                r["Descripción artículo"] = desc or code
                r["Cód. Lista"] = list_code
                r["Desc. Lista"] = list_desc
                r["Moneda"] = default_currency
                r["Unidad"] = "Un"
                all_rows.append(r)

            # Second pass: extract from extra column pairs (multi-column layouts, BUG-6)
            for extra_code_col, extra_price_col in extra_pairs:
                for row in norm[data_start:]:
                    if len(row) <= extra_code_col:
                        continue
                    code = str(row[extra_code_col] if extra_code_col < len(row) else "").strip().upper()
                    if not code or not _is_valid_product_code(code):
                        continue
                    price_str = ""
                    if 0 <= extra_price_col < len(row):
                        price_str = _normalize_price_token(str(row[extra_price_col] or ""))
                        try:
                            if float(price_str) <= 3:
                                price_str = ""
                        except Exception:
                            price_str = ""
                    if not price_str:
                        continue
                    r = _empty_template_row()
                    r["Cód. Artículo"] = code
                    r["Precio"] = price_str
                    r["Descripción artículo"] = code
                    r["Cód. Lista"] = list_code
                    r["Desc. Lista"] = list_desc
                    r["Moneda"] = default_currency
                    r["Unidad"] = "Un"
                    all_rows.append(r)

    deduped: list[dict] = []
    seen: set = set()
    for row in all_rows:
        key = _row_key(row)
        if key in seen or not any(key):
            continue
        seen.add(key)
        deduped.append(row)
    return deduped


# ─────────────────────────────────────────────
# AGENT 1: File Extraction Agent
# ─────────────────────────────────────────────
async def agent_extractor(file_bytes: bytes, filename: str, file_type: str) -> dict:
    """
    Extracts raw text/data from the uploaded file.
    Handles: PDF, XLS/XLSX, CSV, images (JPG, PNG).
    Returns a dict with raw_text and metadata.
    """
    ext = detect_extension(filename, file_bytes)
    raw_text = ""
    extraction_method = "unknown"
    structured_rows = []
    pdf_raw_tables: list = []
    pdf_word_rows: list = []
    metadata = {"filename": filename, "type": ext, "pages": 0}

    def normalize_df_to_records(df: pd.DataFrame) -> list[dict]:
        if df is None or df.empty:
            return []
        local_df = df.copy().fillna("")
        local_df.columns = [str(c).strip() for c in local_df.columns]
        records = local_df.to_dict(orient="records")
        normalized = []
        for rec in records:
            row = {str(k).strip(): str(v).strip() if v is not None else "" for k, v in rec.items()}
            normalized.append(row)
        return normalized

    try:
        if ext in [".pdf"]:
            sync_result = await asyncio.to_thread(_extract_pdf_sync, file_bytes)
            raw_text = sync_result["raw_text"]
            pdf_raw_tables = sync_result["pdf_raw_tables"]
            pdf_word_rows = sync_result["pdf_word_rows"]
            metadata["pages"] = sync_result["pages"]
            metadata["pdf_sources"] = {
                "markitdown_chars": sync_result["md_chars"],
                "pdfplumber_chars": sync_result["plumber_chars"],
            }
            if not raw_text:
                raise HTTPException(status_code=500, detail="PDF extraction produced empty text")
            extraction_method = sync_result["extraction_method"]
            if not metadata["pages"]:
                metadata["pages"] = max(raw_text.count("=== PAGE "), 1)

        elif ext in [".xlsx", ".xlsm"]:
            xl = pd.ExcelFile(io.BytesIO(file_bytes))
            sheets_text = []
            for sheet in xl.sheet_names:
                df = pd.read_excel(io.BytesIO(file_bytes), sheet_name=sheet)
                structured_rows.extend(normalize_df_to_records(df))
                sheets_text.append(f"=== SHEET: {sheet} ===\n{df.to_csv(index=False)}")
            raw_text = "\n".join(sheets_text)

        elif ext in [".xls"]:
            xl = pd.ExcelFile(io.BytesIO(file_bytes), engine="xlrd")
            sheets_text = []
            for sheet in xl.sheet_names:
                df = pd.read_excel(
                    io.BytesIO(file_bytes), engine="xlrd", sheet_name=sheet
                )
                structured_rows.extend(normalize_df_to_records(df))
                sheets_text.append(f"=== SHEET: {sheet} ===\n{df.to_csv(index=False)}")
            raw_text = "\n".join(sheets_text)

        elif ext in [".csv"]:
            # Auto-detect separator (;  ,  \t)
            sample = file_bytes[:4096].decode("utf-8", errors="replace")
            if sample.count(";") > sample.count(","):
                sep = ";"
            elif sample.count("\t") > sample.count(","):
                sep = "\t"
            else:
                sep = ","
            csv_text = file_bytes.decode("utf-8", errors="replace")
            df = pd.read_csv(io.StringIO(csv_text), sep=sep)
            structured_rows.extend(normalize_df_to_records(df))
            raw_text = df.to_csv(index=False, sep=",")

        elif ext in [".docx", ".doc"]:
            doc = DocxDocument(io.BytesIO(file_bytes))
            parts = []
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    parts.append(para.text)
            # Extract tables
            for i, table in enumerate(doc.tables):
                parts.append(f"=== TABLE {i+1} ===")
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    parts.append(" | ".join(cells))
            raw_text = "\n".join(parts)

        elif ext in [".jpg", ".jpeg", ".png", ".webp", ".bmp", ".tiff"]:
            # Use Claude Vision for image files
            img_b64 = base64.standard_b64encode(file_bytes).decode("utf-8")
            media_type = "image/jpeg"
            if ext == ".png":
                media_type = "image/png"
            elif ext == ".webp":
                media_type = "image/webp"
            
            messages = [
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "image",
                            "source": {
                                "type": "base64",
                                "media_type": media_type,
                                "data": img_b64,
                            },
                        },
                        {
                            "type": "text",
                            "text": (
                                "Extract ALL price/product information from this image. "
                                "Include: product codes, descriptions, prices, units, "
                                "currency, discounts, validity dates, and any other "
                                "relevant data. Format as structured text preserving "
                                "all values exactly as shown."
                            ),
                        },
                    ],
                }
            ]
            raw_text = await claude_chat(messages)
            metadata["type"] = "image"

        else:
            raise HTTPException(
                status_code=400, detail=f"Unsupported file type: {ext}"
            )

    except Exception as e:
        raise HTTPException(
            status_code=500, detail=f"Extraction error: {str(e)}"
        )

    if extraction_method == "unknown":
        extraction_method = "vision_ai" if metadata.get("type") == "image" else (
            "direct_structured" if ext in [".csv", ".xlsx", ".xlsm", ".xls"] else "pdf_text_ai"
        )

    return {
        "raw_text": raw_text,
        "structured_rows": structured_rows,
        "pdf_raw_tables": pdf_raw_tables,
        "pdf_word_rows": pdf_word_rows,
        "metadata": metadata,
        "char_count": len(raw_text),
        "extraction_method": extraction_method,
    }


def _norm_hybrid_price(val: object) -> str:
    """Normalize a price value returned by Claude in the hybrid pass."""
    s = re.sub(r"[^\d.,]", "", str(val or ""))
    if re.search(r"\d\.\d{3},\d", s):  # Argentine format 1.234,56
        s = s.replace(".", "").replace(",", ".")
    else:
        s = s.replace(",", ".")
    try:
        return str(round(float(s), 2))
    except Exception:
        return ""


async def _run_hybrid_pass(
    raw_data: dict,
    list_code: str,
    list_desc: str,
    default_currency: str,
) -> list[dict]:
    """Per-page (PDF) or per-chunk (XLS/CSV) Claude pass. Returns normalized rows."""
    file_type = raw_data.get("metadata", {}).get("type", "")
    is_pdf = file_type == ".pdf"

    segments: list[tuple[str, str]] = []  # (text, label)

    if is_pdf:
        for idx, page_text in enumerate(raw_data.get("pdf_pages", []), 1):
            clean = page_text.strip()
            if clean:
                segments.append((clean, f"page {idx}"))
    else:
        raw_text = raw_data.get("raw_text", "").strip()
        if raw_text:
            lines = raw_text.splitlines()
            chunk_lines: list[str] = []
            chars = 0
            part = 1
            for line in lines:
                chunk_lines.append(line)
                chars += len(line) + 1
                if chars >= HYBRID_XLS_CHUNK_CHARS:
                    segments.append(("\n".join(chunk_lines), f"chunk {part}"))
                    chunk_lines, chars, part = [], 0, part + 1
            if chunk_lines:
                segments.append(("\n".join(chunk_lines), f"chunk {part}"))

    if not segments:
        return []

    all_rows: list[dict] = []
    for text, label in segments:
        try:
            raw_response = await claude_chat(
                messages=[{"role": "user", "content": f"{label}:\n{text}"}],
                system=_HYBRID_SYSTEM_PROMPT,
                max_tokens=HYBRID_MAX_TOKENS,
            )
        except Exception:
            continue

        raw_response = raw_response.strip()
        if raw_response.startswith("```"):
            raw_response = raw_response[raw_response.find("\n") + 1:]
            raw_response = re.sub(r"```\s*$", "", raw_response).strip()

        parsed: list[dict] = []
        try:
            parsed = json.loads(raw_response)
            if not isinstance(parsed, list):
                parsed = []
        except Exception:
            # Salvage truncated JSON array
            if "[" in raw_response:
                last_close = raw_response.rfind("}")
                if last_close > 0:
                    try:
                        salvaged = json.loads(raw_response[:last_close + 1] + "]")
                        if isinstance(salvaged, list):
                            parsed = salvaged
                    except Exception:
                        pass

        for item in parsed:
            if not isinstance(item, dict):
                continue
            code = str(item.get("code", "")).strip().upper()
            desc = str(item.get("desc", "")).strip()
            price = _norm_hybrid_price(item.get("price"))
            if not code or not price:
                continue
            r = _empty_template_row()
            r["Cód. Artículo"] = code
            r["Descripción artículo"] = desc
            r["Precio"] = price
            r["Cód. Lista"] = list_code
            r["Desc. Lista"] = list_desc
            r["Moneda"] = default_currency
            r["Unidad"] = "Un"
            all_rows.append(r)

    return all_rows


# ─────────────────────────────────────────────
# AGENT 2: Transformation Agent
# Maps raw text → template columns
# ─────────────────────────────────────────────
async def agent_transformer(
    raw_data: dict,
    supplier_cuit: str = "",
    validation_mode: bool = False,
) -> dict:
    """
    Uses Claude to map raw extracted text to the template columns.
    Returns a dict with 'rows' and 'column_mapping'.
    """
    system_prompt = (
        "Extract product rows from price-list text and return JSON only. "
        "Use exactly two top-level keys: column_mapping and rows. "
        "Each row must include all template columns; use empty string if unknown. "
        "Normalize prices to decimal number text without currency symbols. "
        "Infer currency from context (U$S/USD->USD, EUR->EUR, else ARS). "
        "Use DD/MM/YYYY when dates are present. "
        "If list code/name appears in headers, propagate it to rows."
    )

    def normalize_key(text: str) -> str:
        text = unicodedata.normalize("NFKD", str(text).lower())
        text = "".join(ch for ch in text if not unicodedata.combining(ch))
        text = re.sub(r"[^a-z0-9]+", " ", text)
        return " ".join(text.split())

    def normalize_json_text(raw_text: str) -> str:
        cleaned = raw_text.strip()
        if cleaned.startswith("```"):
            parts = cleaned.split("```")
            if len(parts) > 1:
                cleaned = parts[1]
                if cleaned.startswith("json"):
                    cleaned = cleaned[4:]
        cleaned = cleaned.strip()

        start = cleaned.find("{")
        end = cleaned.rfind("}")
        if start != -1 and end != -1 and end > start:
            cleaned = cleaned[start:end + 1]

        return cleaned.strip()

    def extract_balanced_object(src: str, open_idx: int) -> tuple[str, int]:
        depth = 0
        in_string = False
        escaped = False
        i = open_idx
        while i < len(src):
            ch = src[i]
            if in_string:
                if escaped:
                    escaped = False
                elif ch == "\\":
                    escaped = True
                elif ch == '"':
                    in_string = False
            else:
                if ch == '"':
                    in_string = True
                elif ch == "{":
                    depth += 1
                elif ch == "}":
                    depth -= 1
                    if depth == 0:
                        return src[open_idx:i + 1], i
            i += 1
        return "", -1

    def salvage_rows_from_malformed_json(raw_text: str) -> tuple[dict, list[dict]]:
        column_mapping = {}
        rows = []

        cm_key = '"column_mapping"'
        cm_pos = raw_text.find(cm_key)
        if cm_pos != -1:
            cm_open = raw_text.find("{", cm_pos)
            if cm_open != -1:
                cm_obj, _ = extract_balanced_object(raw_text, cm_open)
                if cm_obj:
                    try:
                        parsed_cm = json.loads(cm_obj)
                        if isinstance(parsed_cm, dict):
                            column_mapping = parsed_cm
                    except json.JSONDecodeError:
                        column_mapping = {}

        rows_key = '"rows"'
        rows_pos = raw_text.find(rows_key)
        if rows_pos == -1:
            return column_mapping, rows

        arr_start = raw_text.find("[", rows_pos)
        if arr_start == -1:
            return column_mapping, rows

        i = arr_start + 1
        while i < len(raw_text):
            ch = raw_text[i]
            if ch == "]":
                break
            if ch != "{":
                i += 1
                continue

            obj_text, obj_end = extract_balanced_object(raw_text, i)
            if not obj_text or obj_end == -1:
                break

            try:
                row = json.loads(obj_text)
                if isinstance(row, dict):
                    rows.append(row)
            except json.JSONDecodeError:
                pass

            i = obj_end + 1

        return column_mapping, rows

    def normalize_rows(rows: list[dict]) -> list[dict]:
        normalized = []
        for row in rows:
            if not isinstance(row, dict):
                continue
            for col in TEMPLATE_COLUMNS:
                if col not in row:
                    row[col] = ""
            normalized.append(row)
        return normalized

    def parse_transform_response(raw_text: str) -> tuple[dict, list[dict]]:
        parsed = json.loads(raw_text)
        column_mapping = {}
        rows = []

        if isinstance(parsed, dict) and "rows" in parsed:
            column_mapping = parsed.get("column_mapping", {})
            rows = parsed["rows"]
        elif isinstance(parsed, list):
            rows = parsed
        else:
            rows = [parsed]

        if not isinstance(column_mapping, dict):
            column_mapping = {}

        return column_mapping, normalize_rows(rows)

    def pick_column(columns: list[str], patterns: list[str]) -> Optional[str]:
        normalized = {c: normalize_key(c) for c in columns}
        for pat in patterns:
            for col, key in normalized.items():
                if pat in key:
                    return col
        return None

    def transform_structured_rows(structured_rows: list[dict]) -> dict:
        if not structured_rows:
            return {"rows": [], "column_mapping": {}}

        columns = list(structured_rows[0].keys())
        code_col = pick_column(columns, ["partid", "part id", "cod art", "sku", "item"])
        if code_col is None:
            # Match "codigo"/"articulo" only when not "codigo de barras" or EAN (BUG-5)
            code_col = next(
                (c for c in columns
                 if any(p in normalize_key(c) for p in ["articulo", "codigo"])
                 and "barras" not in normalize_key(c) and "ean" not in normalize_key(c)),
                None,
            )
        desc_col = pick_column(columns, ["descripcion", "producto", "detalle", "articulo", "nombre"])
        add_col = pick_column(columns, ["adicional", "familia", "rubro", "linea", "modelo", "detalle2"])
        syn_col = pick_column(columns, ["sinonimo", "marca", "alias", "barras", "ean", "barcode"])
        list_col = pick_column(columns, ["cod lista", "lista", "list code"])
        list_desc_col = pick_column(columns, ["desc lista", "lista desc", "list name"])
        currency_col = pick_column(columns, ["moneda", "currency", "divisa"])
        unit_col = pick_column(columns, ["unidad", "um", "u m", "stock um", "presentacion", "medida"])
        price_col = pick_column(columns, ["precio", "unit price", "price", "valor", "importe", "neto"])
        bonif_col = pick_column(columns, ["bonif", "descuento", "dto"])
        from_col = pick_column(columns, ["vigencia desde", "fecha desde", "inicio", "desde"])
        to_col = pick_column(columns, ["vigencia hasta", "fecha hasta", "hasta", "fin"])

        selected = {
            code_col: "Cód. Artículo",
            desc_col: "Descripción artículo",
            add_col: "Descripción adicional artículo",
            syn_col: "Sinónimo",
            list_col: "Cód. Lista",
            list_desc_col: "Desc. Lista",
            currency_col: "Moneda",
            unit_col: "Unidad",
            price_col: "Precio",
            bonif_col: "Bonif.",
            from_col: "Fecha vigencia desde",
            to_col: "Fecha vigencia hasta",
        }
        column_mapping = {src: dst for src, dst in selected.items() if src}

        def get_value(row: dict, col: Optional[str]) -> str:
            if not col:
                return ""
            return str(row.get(col, "") or "").strip()

        out_rows = []
        for row in structured_rows:
            mapped = {
                "Cód. Artículo": get_value(row, code_col),
                "Descripción artículo": get_value(row, desc_col),
                "Descripción adicional artículo": get_value(row, add_col),
                "Sinónimo": get_value(row, syn_col),
                "Cód. Lista": get_value(row, list_col),
                "Desc. Lista": get_value(row, list_desc_col),
                "Moneda": get_value(row, currency_col) or "ARS",
                "Unidad": get_value(row, unit_col),
                "Precio": get_value(row, price_col),
                "Bonif.": get_value(row, bonif_col),
                "Fecha vigencia desde": get_value(row, from_col),
                "Fecha vigencia hasta": get_value(row, to_col),
            }
            if any(v for v in mapped.values()):
                out_rows.append(mapped)

        return {"rows": normalize_rows(out_rows), "column_mapping": column_mapping}

    async def transform_chunk(chunk_text: str) -> tuple[dict, list[dict]]:
        cache_key = hashlib.sha256(f"{CLAUDE_MODEL}|{chunk_text}".encode("utf-8", errors="ignore")).hexdigest()
        cached = _cache_get(cache_key)
        if cached:
            return dict(cached.get("column_mapping", {})), [dict(r) for r in cached.get("rows", [])]

        user_prompt = (
            f"file={raw_data['metadata']['filename']} chars={len(chunk_text)}\n"
            f"columns={json.dumps(TEMPLATE_COLUMNS, ensure_ascii=False)}\n"
            "text:\n"
            f"{chunk_text}"
        )

        try:
            text = await claude_chat(
                messages=[{"role": "user", "content": user_prompt}],
                system=system_prompt,
                max_tokens=6000,
            )
        except HTTPException:
            # Graceful degradation: caller can still use local heuristic rows.
            return {}, []

        text = normalize_json_text(text)

        try:
            parsed_mapping, parsed_rows = parse_transform_response(text)
        except json.JSONDecodeError as e:
            salvaged_mapping, salvaged_rows = salvage_rows_from_malformed_json(text)
            salvaged_rows = normalize_rows(salvaged_rows)
            if salvaged_rows:
                parsed_mapping, parsed_rows = salvaged_mapping, salvaged_rows
            else:
                raise HTTPException(
                    status_code=500,
                    detail=f"Transform parse error: {str(e)}\\nRaw: {text[:500]}",
                )

        _cache_set(cache_key, {"column_mapping": parsed_mapping, "rows": parsed_rows})
        return parsed_mapping, parsed_rows

    structured_rows = raw_data.get("structured_rows") or []
    if structured_rows and raw_data.get("metadata", {}).get("type") in {".csv", ".xlsx", ".xls", ".xlsm"}:
        result = transform_structured_rows(structured_rows)
        mapping_vals = set(result.get("column_mapping", {}).values())
        # Fall through to text/Claude path when column identification fails (BUG-1)
        if "Cód. Artículo" in mapping_vals or "Precio" in mapping_vals:
            return result

    raw_text = raw_data.get("raw_text", "")
    if not raw_text.strip():
        return {"rows": [], "column_mapping": {}}

    list_code, list_desc = _extract_list_header(raw_text)
    default_currency = _infer_currency(raw_text)

    # Detect numeric-item profile (common in some supplier lists like LCT).
    numeric_profile_hits = 0
    for ln in raw_text.splitlines()[:2500]:
        if NUMERIC_ITEM_PROFILE_RE.search(ln or ""):
            numeric_profile_hits += 1
    strict_numeric_profile = (
        raw_data.get("metadata", {}).get("type") == ".pdf"
        and numeric_profile_hits >= 20
    )

    strict_rows_seed = []
    if strict_numeric_profile:
        # Keep strong numeric rows, but don't short-circuit: we still want to recover
        # valid alphanumeric products from the generic path.
        strict_rows_seed = heuristic_extract_rows_numeric_profile(
            raw_text,
            list_code=list_code,
            list_desc=list_desc,
            default_currency=default_currency,
        )

    chunks = build_candidate_blocks(raw_text, max_chunk_chars=6000)
    if not chunks:
        chunks = [raw_text[:6000]]

    all_rows = []
    local_rows_all = []
    ai_rows_all = []
    merged_mapping = {}

    if strict_rows_seed:
        all_rows.extend(strict_rows_seed)
        local_rows_all.extend(strict_rows_seed)
        merged_mapping.update(
            {
                "numeric_profile_code": "Cód. Artículo",
                "numeric_profile_price": "Precio",
                "numeric_profile_desc": "Descripción artículo",
            }
        )

    # Structured table extraction — primary method for complex multi-column layouts.
    # This fixes: missing rows (Termi-Plast, UCA, FL...) and corrupted codes (LY-10→136261).
    pdf_raw_tables = raw_data.get("pdf_raw_tables") or []
    if pdf_raw_tables and raw_data.get("metadata", {}).get("type") == ".pdf":
        table_rows = extract_rows_from_pdf_tables(
            pdf_raw_tables,
            list_code=list_code,
            list_desc=list_desc,
            default_currency=default_currency,
        )
        if table_rows:
            all_rows.extend(table_rows)
            local_rows_all.extend(table_rows)
            merged_mapping["pdf_table_code"] = "Cód. Artículo"
            merged_mapping["pdf_table_price"] = "Precio"

    # Word-coordinate rows supplement table detection for missed multi-column products (BUG-3)
    pdf_word_rows = raw_data.get("pdf_word_rows") or []
    if pdf_word_rows and raw_data.get("metadata", {}).get("type") == ".pdf":
        for row in pdf_word_rows:
            row["Cód. Lista"] = row.get("Cód. Lista") or list_code
            row["Desc. Lista"] = row.get("Desc. Lista") or list_desc
            row["Moneda"] = row.get("Moneda") or default_currency
        all_rows.extend(pdf_word_rows)
        local_rows_all.extend(pdf_word_rows)
        merged_mapping["pdf_word_code"] = "Cód. Artículo"
        merged_mapping["pdf_word_price"] = "Precio"

    ai_enabled = bool(os.getenv("ANTHROPIC_API_KEY", "").strip())
    ai_errors = 0

    if validation_mode:
        # In validation mode we run both engines on all chunks and compare coverage.
        for chunk in chunks:
            local_rows = heuristic_extract_rows(
                chunk,
                list_code=list_code,
                list_desc=list_desc,
                default_currency=default_currency,
            )
            if len(local_rows) < 5:
                local_rows = heuristic_extract_rows_blockwise(
                    chunk,
                    list_code=list_code,
                    list_desc=list_desc,
                    default_currency=default_currency,
                )

            if strict_numeric_profile:
                # In strict profile, only add extra rows that are alphanumeric codes
                # to avoid introducing numeric noise while recovering missing products.
                local_rows = [
                    row for row in local_rows
                    if _is_alphanumeric_code(row.get("Cód. Artículo", ""))
                ]
            local_rows_all.extend(local_rows)

            if ai_enabled:
                chunk_mapping, chunk_rows = await transform_chunk(chunk)
                if not chunk_rows:
                    ai_errors += 1
                merged_mapping.update(chunk_mapping)
                ai_rows_all.extend(chunk_rows)

        all_rows.extend(local_rows_all)
        all_rows.extend(ai_rows_all)
    else:
        ai_chunks = []
        for chunk in chunks:
            local_rows = heuristic_extract_rows(
                chunk,
                list_code=list_code,
                list_desc=list_desc,
                default_currency=default_currency,
            )
            if len(local_rows) < 5:
                local_rows = heuristic_extract_rows_blockwise(
                    chunk,
                    list_code=list_code,
                    list_desc=list_desc,
                    default_currency=default_currency,
                )

            if strict_numeric_profile:
                local_rows = [
                    row for row in local_rows
                    if _is_alphanumeric_code(row.get("Cód. Artículo", ""))
                ]
            local_rows_all.extend(local_rows)

            all_rows.extend(local_rows)

            if not ai_enabled:
                continue

            # Skip AI when numeric-profile seed already covers the list (BUG-7 / LCT case)
            if strict_numeric_profile and len(strict_rows_seed) >= PDF_MIN_EXPECTED_ROWS:
                continue

            # Complement with AI on all chunks by default to maximize recall.
            if AI_COMPLEMENT_ALL_CHUNKS:
                ai_chunks.append(chunk)
            elif len(local_rows) < 8:
                ai_chunks.append(chunk)

        # Hybrid mode: call IA only on weak/non-parseable chunks.
        for chunk in ai_chunks:
            chunk_mapping, chunk_rows = await transform_chunk(chunk)
            if not chunk_rows:
                ai_errors += 1
            merged_mapping.update(chunk_mapping)
            ai_rows_all.extend(chunk_rows)
            all_rows.extend(chunk_rows)

    if all_rows and not merged_mapping:
        merged_mapping = {
            "heuristic_code": "Cód. Artículo",
            "heuristic_price": "Precio",
            "heuristic_list_code": "Cód. Lista",
            "heuristic_list_desc": "Desc. Lista",
        }

    deduped_by_key = {}
    for row in all_rows:
        normalized_row = row if isinstance(row, dict) else {}
        key = _row_key(normalized_row)
        if not key[0] and not key[1]:
            continue

        prev = deduped_by_key.get(key)
        if prev is None:
            deduped_by_key[key] = normalized_row
            continue

        prev_score = _row_non_empty_score(prev)
        curr_score = _row_non_empty_score(normalized_row)
        if curr_score > prev_score:
            deduped_by_key[key] = normalized_row

    deduped = list(deduped_by_key.values())

    # --- Hybrid pass: per-page/chunk Claude extraction merged into deduped base ---
    if ai_enabled and HYBRID_EXTRACTION and not validation_mode:
        hybrid_rows = await _run_hybrid_pass(raw_data, list_code, list_desc, default_currency)
        if hybrid_rows:
            deduped_by_code: dict[str, dict] = {
                str(r.get("Cód. Artículo", "")).strip().upper(): r
                for r in deduped
                if str(r.get("Cód. Artículo", "")).strip()
            }
            for hr in hybrid_rows:
                key = str(hr.get("Cód. Artículo", "")).strip().upper()
                if not key:
                    continue
                if key not in deduped_by_code:
                    deduped_by_code[key] = hr
                else:
                    existing_desc = str(deduped_by_code[key].get("Descripción artículo", "")).strip()
                    new_desc = str(hr.get("Descripción artículo", "")).strip()
                    if len(new_desc) > len(existing_desc):
                        deduped_by_code[key]["Descripción artículo"] = new_desc
            deduped = list(deduped_by_code.values())
            merged_mapping["hybrid_code"] = "Cód. Artículo"
            merged_mapping["hybrid_price"] = "Precio"

    # Second pass: targeted removal of ghost rows in PDFs (BUG-8 + Sonnet audit).
    # Three rules, applied only to PDFs:
    #   R1 – duplicate code + $ in desc → fragmented price artifact from word-coord
    #   R2 – duplicate code + desc==code → word-coord fallback (real desc exists elsewhere)
    #   R3 – model-name code already captured under its numeric code (LCT tool section)
    if raw_data.get("metadata", {}).get("type") == ".pdf":
        _price_in_desc_re = re.compile(r'\$')
        # Count per code and track which codes have at least one real description
        _code_dup_count: dict[str, int] = {}
        _code_has_real_desc: set[str] = set()
        for _r in deduped:
            _c = str(_r.get("Cód. Artículo", "")).strip()
            _d = str(_r.get("Descripción artículo", "")).strip()
            if _c:
                _code_dup_count[_c] = _code_dup_count.get(_c, 0) + 1
                if _d and _d.lower() != _c.lower():
                    _code_has_real_desc.add(_c)
        # Build set of numeric-code rows (code is all digits) to detect model-name ghosts
        _numeric_desc_prefixes: set[str] = set()
        for _r in deduped:
            _c = str(_r.get("Cód. Artículo", "")).strip()
            _d = str(_r.get("Descripción artículo", "")).strip()
            if _c.isdigit():
                # Store the first token of the description as a potential model-name prefix
                _first_token = _d.split()[0] if _d else ""
                if _first_token:
                    _numeric_desc_prefixes.add(_first_token.upper())

        def _is_pdf_ghost(r: dict) -> bool:
            code = str(r.get("Cód. Artículo", "")).strip()
            desc = str(r.get("Descripción artículo", "")).strip()
            n = _code_dup_count.get(code, 0)
            # R1: duplicate code + $ in desc that is a price fragment (< 5 alpha chars).
            # Conservative: skip rows whose desc has real text (paired-table descriptions
            # like "SCA 10 10 5/16 2200 UCA 10 10 45 $ 708,30" have many alpha chars).
            if n > 1 and _price_in_desc_re.search(desc):
                alpha_chars = sum(c.isalpha() for c in desc)
                if alpha_chars < 5:
                    return True
            # R2: duplicate code + desc==code (word-coord fallback, real desc exists)
            if n > 1 and desc.lower() == code.lower() and code in _code_has_real_desc:
                return True
            # R3: non-numeric code whose value appears as first word of a numeric-code row's desc
            if not code.isdigit() and code.upper() in _numeric_desc_prefixes:
                return True
            return False

        deduped = [r for r in deduped if not _is_pdf_ghost(r)]

    recovery_applied = False
    recovery_added = 0
    # Adaptive recovery for difficult PDFs with low initial recall.
    if (
        raw_data.get("metadata", {}).get("type") == ".pdf"
        and ai_enabled
        and len(deduped) < PDF_MIN_EXPECTED_ROWS
    ):
        recovery_chunks = build_full_chunks(raw_text, max_chunk_chars=4500)[:PDF_RECOVERY_MAX_CHUNKS]
        for chunk in recovery_chunks:
            chunk_mapping, chunk_rows = await transform_chunk(chunk)
            if not chunk_rows:
                ai_errors += 1
            merged_mapping.update(chunk_mapping)
            for row in chunk_rows:
                key = _row_key(row)
                if not any(key):
                    continue
                prev = deduped_by_key.get(key)
                if prev is None:
                    deduped_by_key[key] = row
                    recovery_added += 1
                    continue
                if _row_non_empty_score(row) > _row_non_empty_score(prev):
                    deduped_by_key[key] = row
        deduped = list(deduped_by_key.values())
        recovery_applied = True

    response = {
        "rows": normalize_rows(deduped),
        "column_mapping": merged_mapping,
    }

    if ai_errors:
        response["ai_fallback"] = {
            "used": True,
            "failed_chunks": ai_errors,
            "detail": "AI transform failed for one or more chunks; returned local heuristic extraction.",
        }

    if recovery_applied:
        response["adaptive_recovery"] = {
            "applied": True,
            "min_expected_rows": PDF_MIN_EXPECTED_ROWS,
            "added_rows": recovery_added,
            "final_rows": len(deduped),
        }

    if validation_mode:
        local_keys = {_row_key(r) for r in local_rows_all if any(_row_key(r))}
        ai_keys = {_row_key(r) for r in ai_rows_all if any(_row_key(r))}
        final_keys = {_row_key(r) for r in deduped if any(_row_key(r))}

        response["validation"] = {
            "mode": "heuristic_vs_ai",
            "chunks_total": len(chunks),
            "rows_local": len(local_keys),
            "rows_ai": len(ai_keys),
            "rows_final": len(final_keys),
            "missing_in_local_vs_ai": len(ai_keys - local_keys),
            "missing_in_ai_vs_local": len(local_keys - ai_keys),
            "missing_in_final_vs_ai": len(ai_keys - final_keys),
            "missing_in_final_vs_local": len(local_keys - final_keys),
            "examples_missing_in_local_vs_ai": _format_key_samples(ai_keys - local_keys),
            "examples_missing_in_ai_vs_local": _format_key_samples(local_keys - ai_keys),
        }

    return response


# ─────────────────────────────────────────────
# AGENT 3: Verification Agent
# ─────────────────────────────────────────────
async def agent_verifier(rows: list[dict], raw_data: dict) -> dict:
    """
    Validates extracted rows for quality and consistency.
    Returns validation report and cleaned rows.
    """
    issues = []
    cleaned_rows = []

    for i, row in enumerate(rows):
        row_issues = []

        # Check required fields
        if not row.get("Descripción artículo", "").strip():
            row_issues.append("Missing description")

        # Validate price
        price_str = str(row.get("Precio", "")).strip()
        if price_str:
            # Normalize price for both AR and EN number formats.
            price_str = price_str.replace("$", "").replace(" ", "").strip()
            if "," in price_str and "." in price_str:
                if price_str.rfind(",") > price_str.rfind("."):
                    # 1.535,26 -> 1535.26
                    price_str = price_str.replace(".", "").replace(",", ".")
                else:
                    # 1,535.26 -> 1535.26
                    price_str = price_str.replace(",", "")
            elif "," in price_str:
                # 1535,26 -> 1535.26
                price_str = price_str.replace(",", ".")
            try:
                price_val = float(price_str)
                if price_val <= 0:
                    row_issues.append(f"Invalid price: {price_val}")
                else:
                    row["Precio"] = str(round(price_val, 2))
            except ValueError:
                row_issues.append(f"Non-numeric price: '{row.get('Precio')}'")
                row["Precio"] = ""
        else:
            row_issues.append("Missing price")

        # Validate currency
        currency = row.get("Moneda", "").strip().upper()
        if currency and currency not in ["ARS", "USD", "EUR", "US$", "U$S"]:
            row["Moneda"] = "ARS"  # Default
            row_issues.append(f"Normalized currency: {currency} → ARS")

        # Normalize currency aliases
        if row.get("Moneda") in ["US$", "U$S", "USD"]:
            row["Moneda"] = "USD"
        elif row.get("Moneda") == "":
            row["Moneda"] = "ARS"

        # Validate discount
        bonif = str(row.get("Bonif.", "")).strip()
        if bonif:
            try:
                bonif_val = float(bonif.replace(",", ".").replace("%", ""))
                if bonif_val < 0 or bonif_val > 100:
                    row_issues.append(f"Discount out of range: {bonif_val}")
                    row["Bonif."] = ""
                else:
                    row["Bonif."] = str(round(bonif_val, 2))
            except ValueError:
                row["Bonif."] = ""

        if row_issues:
            issues.append({"row": i + 1, "issues": row_issues})

        cleaned_rows.append(row)

    total = len(rows)
    valid = total - len([i for i in issues if any("Missing price" in x or "Missing description" in x for x in i["issues"])])

    report = {
        "total_rows": total,
        "valid_rows": valid,
        "rows_with_issues": len(issues),
        "issues": issues[:20],  # Limit to first 20 for brevity
        "quality_score": round((valid / total * 100) if total > 0 else 0, 1),
    }

    return {"rows": cleaned_rows, "report": report}


# ─────────────────────────────────────────────
# ORCHESTRATOR AGENT
# ─────────────────────────────────────────────
async def orchestrator(
    file_bytes: bytes,
    filename: str,
    supplier_cuit: str = "",
    validation_mode: bool = False,
) -> dict:
    """
    Master agent that coordinates extraction, transformation, and verification.
    """
    log = []

    def log_step(step: str, status: str, detail: str = ""):
        entry = {
            "step": step,
            "status": status,
            "detail": detail,
            "timestamp": datetime.now().isoformat(),
        }
        log.append(entry)
        return entry

    # Step 1: Extract
    log_step("extraction", "running", f"Processing {filename}")
    raw_data = await agent_extractor(file_bytes, filename, "")
    log_step(
        "extraction",
        "done",
        f"Extracted {raw_data['char_count']} chars from {raw_data['metadata'].get('pages', 1)} pages",
    )

    # Step 2: Transform
    log_step("transformation", "running", "Mapping data to template columns")
    transform_result = await agent_transformer(
        raw_data,
        supplier_cuit,
        validation_mode=validation_mode,
    )
    rows = transform_result["rows"]
    column_mapping = transform_result["column_mapping"]
    log_step("transformation", "done", f"Extracted {len(rows)} product rows")

    validation_report = transform_result.get("validation")
    if validation_report:
        log_step(
            "validation",
            "done",
            (
                f"local={validation_report['rows_local']} ai={validation_report['rows_ai']} "
                f"missing_local_vs_ai={validation_report['missing_in_local_vs_ai']}"
            ),
        )

    # Step 3: Verify
    log_step("verification", "running", "Validating and cleaning data")
    result = await agent_verifier(rows, raw_data)
    log_step(
        "verification",
        "done",
        f"Quality: {result['report']['quality_score']}% ({result['report']['valid_rows']}/{result['report']['total_rows']} valid)",
    )

    return {
        "filename": filename,
        "rows": result["rows"],
        "report": result["report"],
        "validation": validation_report,
        "metadata": raw_data["metadata"],
        "extraction_method": raw_data.get("extraction_method", "unknown"),
        "column_mapping": column_mapping,
        "log": log,
    }


# ─────────────────────────────────────────────
# API ENDPOINTS
# ─────────────────────────────────────────────

@app.get("/")
def root():
    return {"status": "ok", "service": "Price Extractor API v1.0"}


@app.post("/extract")
async def extract_file(
    file: UploadFile = File(...),
    supplier_cuit: str = "",
    validation_mode: bool = False,
):
    """
    Main endpoint: upload a price list file and get structured data back.
    Accepts: PDF, XLS, XLSX, CSV, JPG, PNG
    """
    allowed_ext = {".pdf", ".xls", ".xlsx", ".xlsm", ".csv", ".jpg", ".jpeg", ".png", ".webp"}
    file_bytes = await file.read()
    ext = detect_extension(file.filename, file_bytes)
    if ext not in allowed_ext:
        raise HTTPException(status_code=400, detail=f"File type not supported: {ext}")

    if len(file_bytes) > 50 * 1024 * 1024:  # 50MB limit
        raise HTTPException(status_code=400, detail="File too large (max 50MB)")

    result = await orchestrator(
        file_bytes,
        file.filename,
        supplier_cuit,
        validation_mode=validation_mode,
    )
    return JSONResponse(content=result)


@app.post("/extract/download")
async def extract_and_download(
    file: UploadFile = File(...),
    supplier_cuit: str = "",
    format: str = "xlsx",
    validation_mode: bool = False,
):
    """
    Same as /extract but returns a file ready to import.
    format: 'xlsx' or 'xls'
    """
    allowed_ext = {".pdf", ".xls", ".xlsx", ".xlsm", ".csv", ".jpg", ".jpeg", ".png", ".webp"}
    file_bytes = await file.read()
    ext = detect_extension(file.filename, file_bytes)
    if ext not in allowed_ext:
        raise HTTPException(status_code=400, detail=f"File type not supported: {ext}")

    result = await orchestrator(
        file_bytes,
        file.filename,
        supplier_cuit,
        validation_mode=validation_mode,
    )

    cuit_val = supplier_cuit or "30-55555555-1"
    safe_name = Path(file.filename).stem
    df = pd.DataFrame(result["rows"], columns=TEMPLATE_COLUMNS)

    output = io.BytesIO()

    def sanitize_xls_value(value):
        # BIFF8 (.xls) max cell text length is 32767 characters.
        if value is None:
            return ""
        text = str(value)
        if len(text) > 32767:
            text = text[:32767]
        return text

    if format == "xls":
        # Generate .xls (legacy format matching template exactly)
        import xlwt
        try:
            wb = xlwt.Workbook(encoding="utf-8")
            ws = wb.add_sheet("Sheet1")

            # Row 0: CUIT
            ws.write(0, 0, sanitize_xls_value(cuit_val))

            # Row 1: Column headers
            for col_idx, col_name in enumerate(TEMPLATE_COLUMNS):
                ws.write(1, col_idx, sanitize_xls_value(col_name))

            # Row 2+: Data
            for row_idx, row in enumerate(result["rows"]):
                for col_idx, col_name in enumerate(TEMPLATE_COLUMNS):
                    val = sanitize_xls_value(row.get(col_name, ""))
                    ws.write(row_idx + 2, col_idx, val)

            wb.save(output)
        except Exception as e:
            raise HTTPException(
                status_code=500,
                detail=f"XLS generation error: {str(e)}",
            )

        output.seek(0)
        out_filename = f"precios_{safe_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xls"
        media_type = "application/vnd.ms-excel"
    else:
        # Generate .xlsx
        with pd.ExcelWriter(output, engine="openpyxl") as writer:
            header_df = pd.DataFrame([[cuit_val] + [""] * 11])
            header_df.to_excel(writer, index=False, header=False, sheet_name="Sheet1", startrow=0)
            df.to_excel(writer, index=False, sheet_name="Sheet1", startrow=1)

        output.seek(0)
        out_filename = f"precios_{safe_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
        media_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"

    return StreamingResponse(
        output,
        media_type=media_type,
        headers={"Content-Disposition": f"attachment; filename={out_filename}"},
    )


@app.post("/extract/batch")
async def extract_batch(
    files: list[UploadFile] = File(...),
    supplier_cuit: str = "",
    validation_mode: bool = False,
):
    """
    Process multiple files and merge into a single result.
    """
    all_rows = []
    all_logs = []
    all_reports = []

    for file in files:
        file_bytes = await file.read()
        result = await orchestrator(
            file_bytes,
            file.filename,
            supplier_cuit,
            validation_mode=validation_mode,
        )
        all_rows.extend(result["rows"])
        all_logs.extend(result["log"])
        all_reports.append(
            {
                "file": file.filename,
                "report": result["report"],
                "validation": result.get("validation"),
            }
        )

    return JSONResponse(content={
        "total_files": len(files),
        "total_rows": len(all_rows),
        "rows": all_rows,
        "reports": all_reports,
        "log": all_logs,
    })


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000, reload=True)
